"""
Script to convert existing DOCX templates (with real data from the Bangladesh group)
into clean placeholder-based templates for the Workers United web app.

Uses docxtemplater-compatible {tag} format (single curly braces).
"""
import os
import shutil
from docx import Document

# Source templates from the reference folder
SOURCE_DIR = r"C:\VIZE\NEPALCI\13.2.2026"
OUTPUT_DIR = os.path.join(
    r"c:\GitHub Repository for Workers United\Workers-United",
    "public", "templates"
)

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Map of source filename -> output filename
TEMPLATES = {
    "02 UGOVOR O RADU1.docx": "UGOVOR_O_RADU.docx",
    "IZJAVA O SAGLASNOSTI.docx": "IZJAVA_O_SAGLASNOSTI.docx",
    "OVLAŠĆENJE.docx": "OVLASCENJE.docx",
    "POZIVNO PISMO PRUNUS.docx": "POZIVNO_PISMO.docx",
}

# Replacements: old text -> placeholder tag
# Sorted LONGEST FIRST to avoid partial matches
REPLACEMENTS = [
    # Full employer references (longest first)
    ("PRUNUS DOO NOVI SAD, TEMERINSKA 148, NOVI SAD", "{EMPLOYER_FULL_REFERENCE}"),
    ("PRUNUS DOO NOVI SAD (matični broj 21478342)", "{EMPLOYER_NAME} (matični broj {EMPLOYER_MB})"),
    ("PRUNUS DOO NOVI SAD (Company Registration Number: 21478342)", "{EMPLOYER_NAME} (Company Registration Number: {EMPLOYER_MB})"),
    ("PRUNUS DOO NOVI SAD", "{EMPLOYER_NAME}"),
    
    # Full names (longest first)
    ("KHAN MD SAJID", "{WORKER_FULL_NAME}"),
    ("RAHMAN AKIBUR", "{WORKER_FULL_NAME}"),
    ("Rahman Akibur", "{WORKER_FULL_NAME}"),
    
    # First/last names
    ("MD SAJID", "{WORKER_FIRST_NAME}"),
    ("AKIBUR", "{WORKER_LAST_NAME}"),
    ("KHAN", "{WORKER_LAST_NAME}"),
    ("RAHMAN", "{WORKER_LAST_NAME}"),
    
    # Nationality phrases (longest first!)
    ("državljanin Bangladeša", "državljanin {NATIONALITY_SR_GENITIVE}"),
    ("citizen of Bangladeshi", "citizen of {NATIONALITY_EN}"),
    ("BANHLADESHI", "{NATIONALITY_EN}"),
    ("BANGLADESHI", "{NATIONALITY_EN}"),
    ("Bangladeshi", "{NATIONALITY_EN}"),
    ("Bangladshi", "{NATIONALITY_EN}"),
    ("Bangladeša", "{NATIONALITY_SR_GENITIVE}"),
    ("Bangladešu", "{NATIONALITY_SR_LOCATIVE}"),
    ("Bangladesu", "{NATIONALITY_SR_LOCATIVE}"),
    ("Bagladeshi", "{NATIONALITY_EN}"),
    ("In Bangladeshi", "In {NATIONALITY_EN}"),
    
    # Passport
    ("A05976920", "{PASSPORT_NUMBER}"),
    ("A13858776", "{PASSPORT_NUMBER}"),
    
    # DOB
    ("02.03.1988.", "{DATE_OF_BIRTH}"),
    ("02.03.1988", "{DATE_OF_BIRTH}"),
    ("29.01.2001", "{DATE_OF_BIRTH}"),
    
    # Place of birth
    ("FARIDPUR", "{PLACE_OF_BIRTH}"),
    
    # Issuer
    ("DIP/DHAKA.", "{PASSPORT_ISSUER}"),
    ("DIP/DHAKA", "{PASSPORT_ISSUER}"),
    
    # Date of issue
    ("10.12.2022.", "{PASSPORT_ISSUE_DATE}"),
    ("10.12.2022", "{PASSPORT_ISSUE_DATE}"),
    
    # Date of expiry
    ("09.21.2032.", "{PASSPORT_EXPIRY_DATE}"),
    ("09.21.2032", "{PASSPORT_EXPIRY_DATE}"),
    
    # Address
    ("Kashimbad, Kotwali, Kanaipr-7801, Faridpur Bangladeshi", "{WORKER_ADDRESS}"),
    ("Kashimbad, Kotwali, Kanaiptr-7801, Faridpur Bangladeshi", "{WORKER_ADDRESS}"),
    ("Kashimbad, Kotwali, Kanaipur-7801, Faridpur-Bangladshi", "{WORKER_ADDRESS}"),
    ("Kashimbad, Kotwali, Kanaiptr-7801, Faridpur", "{WORKER_ADDRESS}"),
    ("Kashimbad, Kotwali, Kanaipr-7801, Faridpur", "{WORKER_ADDRESS}"),
    ("Kashimbad, Kotwali, Kanaipur-7801, Faridpur", "{WORKER_ADDRESS}"),
    ("Kashimbad, Kotwali, Kanaipur-7801, ", "{WORKER_ADDRESS}"),
    ("Faridpur-Bangladeshi", ""),
    ("Faridpur Bangladeshi", ""),
    ("Faridpur-Bangladshi", ""),
    
    # Employer address variants (longest first)
    ("Temerinska br. 148, 21000 Novi Sad", "{EMPLOYER_ADDRESS}"),
    ("ulica Temerinska 148", "{EMPLOYER_ADDRESS}"),
    ("TEMERINSKA BR. 148", "{EMPLOYER_ADDRESS}"),
    ("Temerinska br. 148", "{EMPLOYER_ADDRESS}"),
    ("TEMERINSKA 148", "{EMPLOYER_ADDRESS}"),
    ("Temerinska 148", "{EMPLOYER_ADDRESS}"),
    
    # Employer details
    ("\"PRUNUS DOO\"", "\"{EMPLOYER_NAME}\""),
    ("PRUNUS DOO", "{EMPLOYER_NAME}"),
    ("PRUNUS", "{EMPLOYER_NAME_SHORT}"),
    ("105549008", "{EMPLOYER_PIB}"),
    ("21478342", "{EMPLOYER_MB}"),
    ("TATJANA CVETKOVIĆ", "{EMPLOYER_DIRECTOR}"),
    ("Tatjana Cvetković", "{EMPLOYER_DIRECTOR}"),
    
    # Contract dates
    ("01.01.2026.", "{CONTRACT_START_DATE}"),
    ("01.01.2026", "{CONTRACT_START_DATE}"),
    ("30.06.2026.", "{CONTRACT_END_DATE}"),
    ("30.06.2026", "{CONTRACT_END_DATE}"),
    
    # Signing dates
    ("29.09.2025.", "{SIGNING_DATE_SR}"),
    ("29.09.2025", "{SIGNING_DATE_SR}"),
    ("September 29, 2025", "{SIGNING_DATE_EN}"),
    
    # Profession
    ("PODOPOLAGAČ", "{JOB_TITLE_SR}"),
    ("PODOPOLAGAC", "{JOB_TITLE_SR}"),
    ("podopolagač", "{JOB_TITLE_SR}"),
    ("Podopolagač", "{JOB_TITLE_SR}"),
    ("Podopolagac", "{JOB_TITLE_SR}"),
    ("Flooring Installer", "{JOB_TITLE_EN}"),
    ("flooring installer", "{JOB_TITLE_EN}"),
    
    # Job description SR (each bullet separately!)
    ("Priprema površina za postavljanje podova, uključujući čišćenje, sušenje i nivelaciju podloge.", "{JOB_DESC_SR_1}"),
    ("Merenje, sečenje i ugradnja različitih vrsta podnih obloga (keramičke pločice, laminat, parket, vinil, tepih, itd.).", "{JOB_DESC_SR_2}"),
    ("Rukovanje alatima i opremom kao što su sekači, valjci, lopatice, lepkovi, nivelišuće", "{JOB_DESC_SR_3}"),
    
    # Job description EN (each bullet separately!)
    ("Preparing surfaces for flooring installation, including cleaning, drying, and leveling the base floor.", "{JOB_DESC_EN_1}"),
    ("Measuring, cutting, and installing various types of floor coverings (ceramic tiles, laminate, parquet, vinyl, carpet, etc.).", "{JOB_DESC_EN_2}"),
    ("Handling tools and equipment such as cutters, rollers, trowels, adhesives, leveling", "{JOB_DESC_EN_3}"),
    
    # Contact
    ("+38166299444", "{CONTACT_PHONE}"),
    ("office@prunus.com", "{CONTACT_EMAIL}"),
    
    # Salary (common value)
    # Note: salary changes per contract, so we leave as is — it's populated from DB
]

# Sort by length of old string, LONGEST FIRST
REPLACEMENTS.sort(key=lambda p: len(p[0]), reverse=True)


def replace_in_paragraph(paragraph):
    """Concatenate all run texts, apply replacements, put back."""
    runs = paragraph.runs
    if not runs:
        return
    
    full_text = ''
    for r in runs:
        full_text += r.text or ''
    
    modified = full_text
    for old, new in REPLACEMENTS:
        if old in modified:
            modified = modified.replace(old, new)
    
    if modified == full_text:
        return
    
    runs[0].text = modified
    for r in runs[1:]:
        r.text = ''


def process_template(src_path, dst_path):
    """Process a single template file."""
    doc = Document(src_path)
    
    # Process paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
    
    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replace_in_paragraph(p)
        if section.footer:
            for p in section.footer.paragraphs:
                replace_in_paragraph(p)
    
    doc.save(dst_path)
    print(f"  Saved: {dst_path}")


def main():
    print("Converting DOCX templates to placeholder format...")
    print(f"Source: {SOURCE_DIR}")
    print(f"Output: {OUTPUT_DIR}\n")
    
    for src_name, dst_name in TEMPLATES.items():
        src_path = os.path.join(SOURCE_DIR, src_name)
        dst_path = os.path.join(OUTPUT_DIR, dst_name)
        
        if not os.path.exists(src_path):
            print(f"  WARNING: {src_path} not found, skipping")
            continue
        
        print(f"Processing: {src_name} -> {dst_name}")
        process_template(src_path, dst_path)
    
    print("\nDone! Templates saved to public/templates/")
    print("Open each in Word to verify placeholder tags are correct.")


if __name__ == "__main__":
    main()
